# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt

doc = Document()

doc.add_heading('化学卡牌闯关 · 游戏介绍', 0)
doc.add_paragraph(
    '《化学卡牌闯关》是一款面向初高中化学学习的网页小游戏。玩家通过"打出一张化学物质卡，'
    '使其与场上的公共卡发生化学反应"来逐步完成关卡目标，在趣味互动中熟悉常见物质的性质与反应类型。'
)

doc.add_heading('一、游戏玩法', level=1)
doc.add_paragraph('本游戏为单人游戏，取消了多人对战与 AI 出牌，提供两种模式：')
doc.add_paragraph('闯关模式：在限定关卡目标下完成各类化学反应，达成目标即通关（无反应提示，需自行判断能否反应）。',
                  style='List Bullet')
doc.add_paragraph('学习模式：自由练习沙盒，无关卡目标、无通关限制，全程开启"反应提示"（拖动手牌时，能反应的公共卡会绿色高亮），'
                  '适合先熟悉物质与反应，再进入闯关挑战。',
                  style='List Bullet')
doc.add_paragraph('两种模式的核心循环一致，如下：')
for line in [
    '1. 开局：系统洗牌后发给玩家 7 张手牌，并在场上铺开 5 张"公共反应卡"。',
    '2. 出牌：用鼠标把一张手牌"拖动"到任意一张能与其发生反应的公共卡上，即可触发反应。'
    '反应成功后，打出的手牌和被反应的公共卡都会进入弃牌堆，并在空位自动补一张新卡（公共区始终保持 5 张）。'
    '成功时还会触发粒子爆炸特效并漂浮显示本次反应类型。',
    '3. 分解：对可分解的卡（双氧水、碳酸钙、氢氧化铜、氢氧化铁、碳酸氢钠）点击右键，'
    '可将其分解为产物并加入手牌。',
    '4. 摸牌："摸牌"按钮常驻（单次最多摸 2 张），手牌上限 10 张；仍无法反应则公共区 5 张全部换新。',
    '5. 通关：每当成功反应一次，系统按反应类型计数；达到本关目标即通关。',
]:
    doc.add_paragraph(line)

doc.add_heading('二、关卡目标', level=1)
doc.add_paragraph('关卡目标在游戏顶部 HUD 实时显示。当前共设置 3 关：')
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = '关卡', '名称', '通关目标'
levels = [
    ('第 1 关', '初出茅庐', '完成 10 次任意化学反应'),
    ('第 2 关', '氧化还原', '完成 15 次反应，其中氧化还原 ≥ 3、置换 ≥ 3'),
    ('第 3 关', '面面俱到', '完成 20 次反应，其中分解 ≥ 2、中和 ≥ 2'),
]
for lv in levels:
    row = table.add_row().cells
    row[0].text, row[1].text, row[2].text = lv
doc.add_paragraph('关卡数据定义在 config.js 的 LEVELS 中，可自由增删或调整目标数值。')

doc.add_heading('三、卡牌与反应', level=1)
doc.add_paragraph('游戏内置 48 种化学物质卡（学习池全量），娱乐池取其高频子集，按类别配色区分：')
doc.add_paragraph('酸（红）、碱（蓝）、盐（紫）、金属单质（黄）、非金属单质（青绿）、氧化物（棕）、条件卡（灰蓝）。',
                  style='Intense Quote')
doc.add_paragraph('反应类型覆盖高中化学主要类别，每次成功反应都会在右侧"历史反应记录"中显示完整的化学方程式与反应类型：')
for line in [
    '化合反应（如 C + O₂ → CO₂）',
    '分解反应（如 2H₂O₂ → 2H₂O + O₂↑）',
    '置换反应（如 Fe + CuSO₄ → FeSO₄ + Cu）',
    '复分解反应（如 HCl + AgNO₃ → AgCl↓ + HNO₃）',
    '中和反应（如 HCl + NaOH → NaCl + H₂O）',
    '氧化还原反应（如 Cu + 4HNO₃ → Cu(NO₃)₂ + 2NO₂↑ + 2H₂O）',
]:
    doc.add_paragraph(line, style='List Bullet')
doc.add_paragraph('完整卡牌清单（含每张卡可反应的卡数、可分解标注）见配套文件《化学卡牌清单.xlsx》。')

doc.add_heading('四、操作说明', level=1)
for line in [
    '出牌：鼠标左键按住手牌，拖到能反应的公共卡上松开（学习模式下，能反应的卡会发绿光提示）。',
    '分解：在可分解卡上点击鼠标右键。',
    '查看反应：每次成功出牌/分解，右侧面板都会记录方程式与类型。',
    '摸牌：点击常驻的"摸牌"按钮（学习模式也可随时摸牌练习）。',
]:
    doc.add_paragraph(line)

doc.add_heading('五、道具系统', level=1)
doc.add_paragraph('游戏内置三种道具，闯关与学习模式均可使用。每进入新关卡，三种道具初始各 1 次；'
                  '之后每完成 3 次成功反应，系统会随机奖励一个道具的使用次数（屏幕顶部会提示"获得道具"）。')
for line in [
    '🗑️ 垃圾桶（拖动使用）：把垃圾桶图标拖到一张手牌上松开，即可将该手牌弃入弃牌堆。',
    '💡 灯泡（点击使用）：点击灯泡，系统会随机高亮一组"场上公共卡 + 可与之反应的手牌"（黄色光晕约 2.5 秒）；'
    '若当前没有任何可反应的组合，则直接刷新全部 5 张场上卡。',
    '🧪 烧杯（拖动使用）：把烧杯图标拖到一张手牌上松开，该手牌与场上随机一张卡交换位置。',
]:
    doc.add_paragraph(line, style='List Bullet')
doc.add_paragraph('道具栏位于手牌区下方，显示每种道具的剩余次数；次数为 0 时图标置灰、不可使用。',
                  style='Intense Quote')

doc.add_heading('六、设置', level=1)
doc.add_paragraph('在"设置"页可调整：')
for line in [
    '背景颜色：内置蓝白（默认）、暗夜、清新绿、暖橙四套主题，全局生效。',
    '音效总开关、背景音乐开关、音量。',
    '卡牌类别配色开关（关闭则所有卡牌统一配色）。',
    '恢复默认按钮可一键还原。',
]:
    doc.add_paragraph(line, style='List Bullet')

doc.add_heading('七、如何运行', level=1)
doc.add_paragraph('本游戏为纯前端、零构建网页，无需安装任何环境。直接用浏览器打开 chem-cards 文件夹下的 '
                  'index.html 即可开始游玩（建议 Chrome / Edge 等现代浏览器）。')
doc.add_paragraph('文件结构：')
for line in [
    'index.html —— 首页与关卡选择',
    'game.html —— 闯关对局（含拖拽出牌、右键分解、历史记录）',
    'cardpool.html —— 牌库展示',
    'setting.html —— 设置',
    'config.js —— 卡池、反应、关卡、主题等全部数据',
    'common.js —— 主题、音效、卡牌渲染等共享逻辑',
    'styles.css —— 界面样式',
]:
    doc.add_paragraph(line, style='List Bullet')

doc.save(r'C:\Users\lkm49\WorkBuddy\2026-08-01-10-11-57\chem-cards\化学卡牌闯关游戏介绍.docx')
print('docx saved')
